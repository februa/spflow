"""整相方式の設計結果と評価結果を統合した日本語Word文書を生成する。"""

# pyright: reportMissingImports=false
# python-docxは製品依存ではなくCodex文書runtimeから供給される成果物生成専用依存である。

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from html import escape
from pathlib import Path
from typing import Any

from latex2mathml.converter import convert as latex_to_mathml
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tools.render_formal_s2a_t2a_word import (
    CAUTION_FILL,
    HEADER_FILL,
    MUTED,
    TABLE_WIDTH_DXA,
    _add_body,
    _add_bullet,
    _add_heading,
    _configure_document,
    _set_cell_margins,
    _set_cell_shading,
    _set_run_font,
    _set_table_geometry,
)


REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MARKDOWN = REPOSITORY_ROOT / "doc" / "SpFlow" / "整相方式設計結果.md"
OUTPUT_PATH = REPOSITORY_ROOT / "output" / "word" / "整相方式設計・評価結果.docx"
FLOW_ASSET_DIR = REPOSITORY_ROOT / "doc" / "SpFlow" / "assets" / "alignment_flows"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MATHML_NAMESPACE = "http://www.w3.org/1998/Math/MathML"
OMML_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
YU_GOTHIC_MEDIUM = Path(
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/"
    "ee89e7987a76cc8cfdff36c96bd7bc77655b343e.asset/AssetData/YuGothic-Medium.otf"
)
YU_GOTHIC_BOLD = Path(
    "/System/Library/AssetsV2/com_apple_MobileAsset_Font8/"
    "b7a6a6575a699e801915b73b9e1e75c74a3404ce.asset/AssetData/YuGothic-Bold.otf"
)


def _clean_inline_markdown(text: str) -> str:
    """Word本文へ不要なMarkdown記号だけを除去し、技術記号は保持する。"""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("__", "")
    return text.replace("`", "")


def _load_flow_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    """日本語flow図をbitmap化するためYu Gothicの実fontを読み込む。"""
    path = YU_GOTHIC_BOLD if bold else YU_GOTHIC_MEDIUM
    if not path.exists():
        raise FileNotFoundError(path)
    return ImageFont.truetype(str(path), size=size)


def _wrap_flow_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    """box幅を越えない位置で日本語・英数字を文字単位に折り返す。"""
    lines: list[str] = []
    current = ""
    for character in text:
        candidate = current + character
        if current and draw.textbbox((0, 0), candidate, font=font)[2] > width:
            lines.append(current)
            current = character
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def _draw_box(
    draw: ImageDraw.ImageDraw,
    position: tuple[int, int, int, int],
    text: str,
    *,
    fill: str,
    font: ImageFont.FreeTypeFont,
) -> None:
    """flow nodeを角丸boxと中央揃えlabelで描画する。"""
    left, top, right, bottom = position
    draw.rounded_rectangle(position, radius=16, fill=fill, outline="#426B8A", width=3)
    lines = _wrap_flow_text(draw, text, font, right - left - 28)
    line_height = font.size + 6
    y = top + ((bottom - top) - line_height * len(lines)) / 2
    for line in lines:
        bounds = draw.textbbox((0, 0), line, font=font)
        x = left + ((right - left) - (bounds[2] - bounds[0])) / 2
        draw.text((x, y), line, font=font, fill="#132638")
        y += line_height


def _draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
) -> None:
    """処理方向を線とarrow headで描画し、必要なら軸labelを付与する。"""
    draw.line((start, end), fill="#2E5D7B", width=4)
    x, y = end
    draw.polygon(((x, y), (x - 14, y - 8), (x - 14, y + 8)), fill="#2E5D7B")
    if label:
        font = _load_flow_font(22, bold=True)
        midpoint = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2 - 30)
        draw.text(midpoint, label, font=font, fill="#2E5D7B", anchor="mm")


def _write_svg_flow(
    path: Path,
    title: str,
    boxes: tuple[tuple[int, int, int, int, str, str], ...],
    arrows: tuple[tuple[int, int, int, int, str], ...],
) -> None:
    """Word外でも再利用できる同内容のSVG sourceを保存する。"""
    parts = [
        '<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="720" viewBox="0 0 1600 720">',
        '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#2E5D7B"/></marker></defs>',
        '<rect width="1600" height="720" fill="white"/>',
        f'<text x="800" y="48" text-anchor="middle" font-family="Yu Gothic" font-size="30" font-weight="bold" fill="#173A55">{escape(title)}</text>',
    ]
    for x1, y1, x2, y2, text_value, fill in boxes:
        parts.append(f'<rect x="{x1}" y="{y1}" width="{x2-x1}" height="{y2-y1}" rx="16" fill="{fill}" stroke="#426B8A" stroke-width="3"/>')
        parts.append(f'<foreignObject x="{x1+12}" y="{y1+10}" width="{x2-x1-24}" height="{y2-y1-20}"><div xmlns="http://www.w3.org/1999/xhtml" style="font-family:Yu Gothic;font-size:23px;text-align:center;color:#132638;display:flex;align-items:center;justify-content:center;height:100%;">{escape(text_value)}</div></foreignObject>')
    for x1, y1, x2, y2, label in arrows:
        parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#2E5D7B" stroke-width="4" marker-end="url(#arrow)"/>')
        if label:
            parts.append(f'<text x="{(x1+x2)//2}" y="{(y1+y2)//2-16}" text-anchor="middle" font-family="Yu Gothic" font-size="22" font-weight="bold" fill="#2E5D7B">{escape(label)}</text>')
    parts.append("</svg>")
    path.write_text("\n".join(parts), encoding="utf-8")


def _generate_flow_assets() -> tuple[Path, ...]:
    """4つのMermaid flowをSVG sourceとWord挿入用PNGへ変換する。"""
    FLOW_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    definitions = (
        (
            "overall_method_flow",
            "整相方式を構成する3軸と重み設計部品",
            (
                (40, 250, 270, 390, "整数遅延前センサ入力 X", "#E8EEF5"),
                (330, 100, 570, 220, "S共分散\n同一時刻block", "#F2F4F7"),
                (330, 420, 570, 540, "T共分散\n候補方位別整合", "#EAF4EA"),
                (650, 250, 900, 390, "EBAE または MVDR\n完成重み設計", "#FFF4CE"),
                (980, 80, 1260, 200, "S1 / T1\n元入力座標の直接FIR", "#E8EEF5"),
                (980, 290, 1260, 410, "整数遅延位相 D で\n残差座標へ回転", "#F2F4F7"),
                (1320, 210, 1560, 330, "S2a / T2a\n残差直接FIR", "#EAF4EA"),
                (1320, 430, 1560, 550, "S2b / T2b\n主枝－差分枝", "#FCE8E6"),
            ),
            (
                (270, 300, 330, 160, "S"), (270, 340, 330, 480, "T"),
                (570, 160, 650, 300, ""), (570, 480, 650, 340, ""),
                (900, 300, 980, 140, "1"), (900, 350, 980, 350, "2"),
                (1260, 350, 1320, 270, "a"), (1260, 370, 1320, 490, "b"),
            ),
        ),
        (
            "covariance_flow",
            "S/T共分散生成フロー",
            (
                (80, 270, 380, 410, "整数遅延前snapshot", "#E8EEF5"),
                (520, 120, 870, 250, "同一時刻block\nS共分散 R_S", "#F2F4F7"),
                (520, 430, 870, 560, "候補方位別の時間・位相整合\nT共分散 R_T(θ_b)", "#EAF4EA"),
                (1040, 270, 1410, 410, "共通の重み設計部品\nEBAE または MVDR", "#FFF4CE"),
            ),
            ((380, 310, 520, 185, "S"), (380, 370, 520, 495, "T"), (870, 185, 1040, 315, ""), (870, 495, 1040, 365, "")),
        ),
        (
            "coordinate_flow",
            "1/2 FIR実現座標フロー",
            (
                (50, 270, 300, 410, "完成共分散 R_C\nsteering a", "#E8EEF5"),
                (430, 100, 760, 230, "方式1\n元入力座標重み w_C", "#F2F4F7"),
                (430, 440, 760, 570, "方式2\nR_C2=D R_C D^H\na_C2=D a", "#EAF4EA"),
                (900, 100, 1240, 230, "元入力 X に直接FIR\nY=w_C^H X", "#E8EEF5"),
                (900, 440, 1240, 570, "整数遅延入力 X_D=D X\n残差重み v_C", "#FFF4CE"),
                (1370, 270, 1570, 410, "同じ完成出力\nD^H v_C=w_C", "#FCE8E6"),
            ),
            ((300, 310, 430, 165, "1"), (300, 370, 430, 505, "2"), (760, 165, 900, 165, ""), (760, 505, 900, 505, ""), (1240, 165, 1370, 315, ""), (1240, 505, 1370, 365, "")),
        ),
        (
            "realization_flow",
            "a/b実現構造フロー",
            (
                (50, 270, 330, 410, "整数遅延後入力 X_D\n完成残差重み v_C", "#E8EEF5"),
                (480, 100, 820, 230, "2a: 残差完成重みを\n1枝で直接FIR化", "#EAF4EA"),
                (480, 430, 820, 560, "差分重み q_C=f-v_C\n固定主枝と差分枝", "#FFF4CE"),
                (980, 100, 1280, 230, "Y_2a=v_C^H X_D", "#E8EEF5"),
                (980, 430, 1280, 560, "主枝出力－差分枝出力", "#FCE8E6"),
                (1400, 270, 1580, 410, "Y_2a = Y_2b", "#EAF4EA"),
            ),
            ((330, 310, 480, 165, "a"), (330, 370, 480, 495, "b"), (820, 165, 980, 165, ""), (820, 495, 980, 495, ""), (1280, 165, 1400, 315, ""), (1280, 495, 1400, 365, "")),
        ),
    )
    generated: list[Path] = []
    title_font = _load_flow_font(30, bold=True)
    node_font = _load_flow_font(23)
    for stem, title, boxes, arrows in definitions:
        svg_path = FLOW_ASSET_DIR / f"{stem}.svg"
        png_path = FLOW_ASSET_DIR / f"{stem}.png"
        _write_svg_flow(svg_path, title, boxes, arrows)
        image = Image.new("RGB", (1600, 720), "white")
        draw = ImageDraw.Draw(image)
        draw.text((800, 42), title, font=title_font, fill="#173A55", anchor="mm")
        for x1, y1, x2, y2, node_text, fill in boxes:
            _draw_box(draw, (x1, y1, x2, y2), node_text, fill=fill, font=node_font)
        for x1, y1, x2, y2, label in arrows:
            _draw_arrow(draw, (x1, y1), (x2, y2), label)
        image.save(png_path, format="PNG", dpi=(160, 160))
        generated.append(png_path)
    return tuple(generated)


def _math_run(text: str) -> Any:
    """MathML tokenをWord数式runへ変換する。"""
    run = OxmlElement("m:r")
    properties = OxmlElement("m:rPr")
    style = OxmlElement("m:sty")
    style.set(qn("m:val"), "p")
    properties.append(style)
    run.append(properties)
    token = OxmlElement("m:t")
    token.text = text
    run.append(token)
    return run


def _append_mathml(parent: Any, element: ET.Element) -> None:
    """latex2mathmlの主要MathML要素をWord OMMLへ再帰変換する。"""
    tag = element.tag.rsplit("}", maxsplit=1)[-1]
    children = list(element)
    if tag in {"math", "mrow", "mstyle", "semantics"}:
        for child in children:
            _append_mathml(parent, child)
        return
    if tag in {"mi", "mn", "mo", "mtext"}:
        parent.append(_math_run(element.text or ""))
        return
    if tag == "mfrac" and len(children) >= 2:
        fraction = OxmlElement("m:f")
        numerator = OxmlElement("m:num")
        denominator = OxmlElement("m:den")
        _append_mathml(numerator, children[0])
        _append_mathml(denominator, children[1])
        fraction.extend((numerator, denominator))
        parent.append(fraction)
        return
    script_map = {"msub": ("m:sSub", "m:sub"), "msup": ("m:sSup", "m:sup")}
    if tag in script_map and len(children) >= 2:
        outer_name, script_name = script_map[tag]
        outer = OxmlElement(outer_name)
        base = OxmlElement("m:e")
        script = OxmlElement(script_name)
        _append_mathml(base, children[0])
        _append_mathml(script, children[1])
        outer.extend((base, script))
        parent.append(outer)
        return
    if tag == "msubsup" and len(children) >= 3:
        outer = OxmlElement("m:sSubSup")
        base, subscript, superscript = OxmlElement("m:e"), OxmlElement("m:sub"), OxmlElement("m:sup")
        _append_mathml(base, children[0])
        _append_mathml(subscript, children[1])
        _append_mathml(superscript, children[2])
        outer.extend((base, subscript, superscript))
        parent.append(outer)
        return
    if tag in {"msqrt", "mroot"} and children:
        radical = OxmlElement("m:rad")
        degree = OxmlElement("m:deg")
        expression = OxmlElement("m:e")
        if tag == "mroot" and len(children) >= 2:
            _append_mathml(degree, children[1])
        _append_mathml(expression, children[0])
        radical.extend((degree, expression))
        parent.append(radical)
        return
    if tag in {"munder", "mover", "munderover"} and children:
        # 積分・総和記号の上下限はscriptとして保持し、Word数式modeで編集可能にする。
        outer = OxmlElement("m:sSubSup")
        base, subscript, superscript = OxmlElement("m:e"), OxmlElement("m:sub"), OxmlElement("m:sup")
        _append_mathml(base, children[0])
        if len(children) >= 2:
            _append_mathml(subscript if tag != "mover" else superscript, children[1])
        if len(children) >= 3:
            _append_mathml(superscript, children[2])
        outer.extend((base, subscript, superscript))
        parent.append(outer)
        return
    if tag == "mtable":
        matrix = OxmlElement("m:m")
        for row_element in children:
            row = OxmlElement("m:mr")
            for cell_element in list(row_element):
                cell = OxmlElement("m:e")
                _append_mathml(cell, cell_element)
                row.append(cell)
            matrix.append(row)
        parent.append(matrix)
        return
    for child in children:
        _append_mathml(parent, child)


def _add_native_equation(document: Document, latex_lines: list[str]) -> None:
    """LaTeX blockをMathML経由でWordネイティブOMML数式へ変換する。"""
    latex = " ".join(line.strip() for line in latex_lines)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    equation_paragraph = OxmlElement("m:oMathPara")
    equation = OxmlElement("m:oMath")
    try:
        mathml = ET.fromstring(latex_to_mathml(latex))
        _append_mathml(equation, mathml)
    except (ET.ParseError, ValueError):
        equation.append(_math_run(latex))
    equation_paragraph.append(equation)
    paragraph._p.append(equation_paragraph)


def _add_title_block(document: Document) -> None:
    """設計と評価を統合した文書であることを冒頭に明示する。"""
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run("整相方式 設計・評価結果"), size_pt=24, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run_font(
        subtitle.add_run("S/T共分散 × 直接/整数遅延分離 × a/b実現方式"),
        size_pt=13,
        color=MUTED,
    )
    _add_body(
        document,
        "本書は整相方式の数式、処理順、方式間の等価性、有限FIR実現、成立条件、"
        "および正式S2a/T2a評価結果を一体としてレビューするWord版である。"
        "EBAEとMVDRは重み設計部品、S/T・1/2・a/bは共分散とFIR実現座標の選択として分離して記述する。",
    )


def _mark_header_row(table: Any) -> None:
    """screen readerと改ページ後の再表示のため先頭行をheaderにする。"""
    properties = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _add_method_flow(document: Document) -> None:
    """方式軸と処理順を図と定義表で欠落なく示す。"""
    _add_heading(document, "0. 方式全体の処理フロー", 1)
    _add_body(
        document,
        "方式名は、共分散構成（S/T）、FIR実現座標（1/2）、整数遅延後の実現構造（a/b）の3軸で読む。"
        "EBAEまたはMVDRは完成周波数重みを生成する交換可能な部品であり、整相方式名には含めない。",
    )

    _add_flow_image(
        document,
        FLOW_ASSET_DIR / "overall_method_flow.png",
        "S/T・1/2・a/b方式の処理順と重み設計部品の責務分離",
        1,
    )

    headers = ("方式", "共分散", "重み設計前の座標変換", "時間領域実現", "意味")
    rows = (
        ("S1", "元入力の同一時刻S共分散", "なし", "元入力座標で完成重みを直接FIR化", "S共分散＋直接FIR"),
        ("S2a", "整数遅延前入力からS共分散", "整数遅延分の位相回転", "実整数遅延buffer＋残差完成重みFIR", "S共分散＋整数遅延分離"),
        ("S2b", "S2aと同じS共分散", "S2aと同じ位相回転", "固定主枝－差分補正枝", "S2aと同じ完成重みの別構造"),
        ("T1", "候補方位別時間切り出しT共分散", "切り出し時刻差の位相を元入力座標へ補正", "元入力座標で完成重みを直接FIR化", "T共分散＋直接FIR"),
        ("T2a", "候補方位別時間切り出しT共分散", "整数遅延後の残差座標へ回転", "実整数遅延buffer＋残差完成重みFIR", "T共分散＋整数遅延分離"),
        ("T2b", "T2aと同じT共分散", "T2aと同じ位相回転", "固定主枝－差分補正枝", "T2aと同じ完成重みの別構造"),
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(text), size_pt=8.0, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _set_run_font(paragraph.add_run(value), size_pt=8.0)
    _set_table_geometry(table, (800, 2050, 2100, 2350, 2060))
    _mark_header_row(table)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(caption.add_run("表0-1　方式定義表"), size_pt=9.0, color=MUTED)

    for text in (
        "S2a/T2aのa方式: 完成残差重みを直接FIR化して整数遅延後信号へ適用する。",
        "S2b/T2bのb方式: 固定整相主枝から差分補正枝を減算し、a方式と同じ完成出力を得る。",
        "Ns=0ではEBAE補正項が存在せず、EBAE重みはCBFと一致する。",
        "S2a/T2aの共分散は整数遅延前入力から生成し、重み側だけを整数遅延後の残差座標へ回転する。",
    ):
        _add_bullet(document, text)


def _add_markdown_table(document: Document, lines: list[str]) -> None:
    """Markdown表を内容幅に応じた固定DXA表へ変換する。"""
    parsed = [[_clean_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(parsed) < 2:
        return
    headers = parsed[0]
    data_rows = parsed[2:] if all(set(cell) <= {"-", ":"} for cell in parsed[1]) else parsed[1:]
    column_count = len(headers)
    if column_count == 0:
        return
    # 短い識別列を狭くし、説明列へ残りを配分する。全cell幅の合計は9360 DXAに固定する。
    weights = [max(3, min(16, len(header))) for header in headers]
    if column_count >= 2:
        weights[-1] = max(weights[-1], 18)
    total_weight = sum(weights)
    widths = [max(600, int(TABLE_WIDTH_DXA * weight / total_weight)) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    if widths[-1] < 600:
        deficit = 600 - widths[-1]
        widths[-1] = 600
        widths[0] -= deficit

    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(text), size_pt=7.8, bold=True)
    for row in data_rows:
        cells = table.add_row().cells
        normalized = row + [""] * (column_count - len(row))
        for index, text in enumerate(normalized[:column_count]):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == column_count - 1 else WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(paragraph.add_run(text), size_pt=7.8)
    _set_table_geometry(table, tuple(widths))
    _mark_header_row(table)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_record_blocks(document: Document, lines: list[str]) -> None:
    """横長のMarkdown表を1行1項目の縦型recordへ変換する。"""
    parsed = [[_clean_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(parsed) < 3:
        return
    headers = parsed[0]
    rows = parsed[2:] if all(set(cell) <= {"-", ":"} for cell in parsed[1]) else parsed[1:]
    for row in rows:
        normalized = row + [""] * (len(headers) - len(row))
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(7)
        heading.paragraph_format.space_after = Pt(2)
        _set_run_font(heading.add_run(normalized[0]), size_pt=10.5, bold=True, color=DARK_BLUE)
        for label, value in zip(headers[1:], normalized[1:]):
            _add_body(document, value, bold_lead=f"{label}: ")


def _add_flow_image(document: Document, path: Path, caption: str, figure_number: int) -> None:
    """SVGと同内容のPNG flow図をWordへ挿入しalt textを付与する。"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    alt_text = f"図{figure_number} {caption}"
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", f"図{figure_number}")
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(caption_paragraph.add_run(alt_text), size_pt=9.0, color=MUTED)


def _add_source_image(document: Document, markdown_path: str, caption: str, figure_number: int) -> None:
    """Markdown基準の相対画像を解決し、本文幅とalt textを設定する。"""
    path = (SOURCE_MARKDOWN.parent / markdown_path).resolve()
    if not path.exists():
        _add_body(document, f"[画像未生成: {path}]", bold_lead="注意: ")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    alt_text = f"図{figure_number} {caption}"
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", f"図{figure_number}")
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    _set_run_font(caption_paragraph.add_run(alt_text), size_pt=9.0, color=MUTED)


def _add_code_or_equation(document: Document, lines: list[str]) -> None:
    """数式とcodeを改行保持のmonospace blockとして表示する。"""
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_cell_shading(table.cell(0, 0), "F7F8FA")
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Menlo"
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    _set_table_geometry(table, (TABLE_WIDTH_DXA,))
    # 1 cellの数式blockもアクセシビリティ監査上の先頭行を明示する。
    _mark_header_row(table)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _render_markdown(document: Document) -> None:
    """既存設計書の全章をWordの見出し、表、図、本文へ変換する。"""
    lines = SOURCE_MARKDOWN.read_text(encoding="utf-8").splitlines()
    index = 0
    figure_number = 2
    mermaid_index = 0
    current_heading = ""
    flow_images = (
        FLOW_ASSET_DIR / "overall_method_flow.png",
        FLOW_ASSET_DIR / "covariance_flow.png",
        FLOW_ASSET_DIR / "coordinate_flow.png",
        FLOW_ASSET_DIR / "realization_flow.png",
    )
    flow_captions = (
        "整相方式を構成する3軸と重み設計部品",
        "S/T共分散生成フロー",
        "1/2 FIR実現座標フロー",
        "a/b実現構造フロー",
    )
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            _add_body(document, _clean_inline_markdown(" ".join(paragraph_buffer)))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped.removeprefix("```").strip().lower()
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            if language == "mermaid":
                # 最初の全体flowは冒頭の図1ですでに示すため重複掲載しない。
                if mermaid_index > 0 and mermaid_index < len(flow_images):
                    _add_flow_image(
                        document,
                        flow_images[mermaid_index],
                        flow_captions[mermaid_index],
                        figure_number,
                    )
                    figure_number += 1
                mermaid_index += 1
            else:
                _add_code_or_equation(document, block)
            index += 1
            continue
        if stripped == "\\[":
            flush_paragraph()
            block = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                block.append(lines[index])
                index += 1
            _add_native_equation(document, block)
            index += 1
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = min(3, len(heading_match.group(1)))
            text = _clean_inline_markdown(heading_match.group(2))
            if level == 1 and text == "整相方式 設計結果":
                index += 1
                continue
            _add_heading(document, text, level)
            current_heading = text
            index += 1
            continue
        image_match = re.match(r"^!\[([^]]*)\]\(([^)]+)\)$", stripped)
        if image_match:
            flush_paragraph()
            _add_source_image(document, image_match.group(2), image_match.group(1), figure_number)
            figure_number += 1
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            if current_heading.startswith("6.4") or current_heading.startswith("13.3"):
                _add_record_blocks(document, table_lines)
            else:
                _add_markdown_table(document, table_lines)
            continue
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            _add_bullet(document, _clean_inline_markdown(bullet_match.group(1)))
            index += 1
            continue
        number_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if number_match:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            _set_run_font(paragraph.add_run(_clean_inline_markdown(number_match.group(2))), size_pt=10.5)
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def build_document() -> Path:
    """設計書全章と評価成果物を統合した日本語DOCXを生成する。"""
    document = Document()
    _configure_document(document)
    _generate_flow_assets()
    document.sections[0].header.paragraphs[0].clear()
    header = document.sections[0].header.paragraphs[0]
    _set_run_font(header.add_run("SpFlow | 整相方式 設計・評価結果"), size_pt=9.0, color=MUTED)
    _add_title_block(document)
    _add_method_flow(document)
    document.add_page_break()
    _render_markdown(document)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
