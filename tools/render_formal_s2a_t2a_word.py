"""正式S2a/T2a成立性評価結果をレビュー用Word文書へ出力する。"""

# pyright: reportMissingImports=false
# python-docxは製品実行依存ではなく、Codex文書runtimeから供給される成果物生成専用依存である。

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = REPOSITORY_ROOT / "output" / "word" / "正式S2a_T2a成立性評価結果.docx"
ARTIFACT_DIR = (
    REPOSITORY_ROOT
    / "artifacts"
    / "beamforming"
    / "formal_s2a_t2a_endfire"
    / "review_pack"
)

# 日本語Word文書の標準font。実機Wordでは日本語glyphを含むYu Gothicを使用する。
FONT_NAME = "Yu Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
HEADER_FILL = "F2F4F7"
CAUTION_FILL = "FFF4CE"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_run_font(
    run: Any,
    *,
    size_pt: float,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    """日本語を含むrunへWordとLibreOfficeで共通のfont指定を設定する。"""
    font = run.font
    font.name = FONT_NAME
    font.size = Pt(size_pt)
    font.bold = bold
    if color is not None:
        font.color.rgb = color
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_cell_shading(cell: Any, fill: str) -> None:
    """headerや注意欄を本文と識別できる淡色で塗る。"""
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any) -> None:
    """表文字が罫線へ密着しないよう上下80、左右120 DXAを確保する。"""
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: tuple[int, ...]) -> None:
    """表全体、grid、各cell幅を同一DXA値へ固定する。"""
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("table column widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440.0)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _add_table(
    document: Document,
    headers: tuple[str, ...],
    rows: tuple[tuple[str, ...], ...],
    widths_dxa: tuple[int, ...],
) -> None:
    """比較可能な数値だけを固定geometryの表として追加する。"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(text), size_pt=8.5, bold=True)
    # 複数ページへ分割された場合も列の意味を失わないよう、先頭行を正式なheader rowにする。
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)
    for values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == len(values) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            )
            _set_run_font(paragraph.add_run(text), size_pt=8.5)
    _set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_heading(document: Document, text: str, level: int) -> None:
    """standard_business_brief presetの見出し階層を追加する。"""
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    size = {1: 16.0, 2: 13.0, 3: 12.0}[level]
    color = BLUE if level < 3 else DARK_BLUE
    _set_run_font(run, size_pt=size, bold=True, color=color)


def _add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    """本文を6 pt後・1.10行間で追加する。"""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead is not None:
        _set_run_font(paragraph.add_run(bold_lead), size_pt=10.5, bold=True)
    _set_run_font(paragraph.add_run(text), size_pt=10.5)


def _add_bullet(document: Document, text: str) -> None:
    """Wordのlist定義を使い、折返し行を本文位置へ揃える。"""
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.10
    _set_run_font(paragraph.add_run(text), size_pt=10.5)


def _add_figure(document: Document, filename: str, caption: str) -> None:
    """評価PNGを本文幅へ収め、図とcaptionを同じ位置に配置する。"""
    path = ARTIFACT_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(6.4))
    # Wordのscreen readerが図の目的を把握できるよう、captionを代替テキストにも設定する。
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption.split(".", maxsplit=1)[0])
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    _set_run_font(caption_paragraph.add_run(caption), size_pt=9.0, color=MUTED)


def _add_page_number(paragraph: Any) -> None:
    """WordのPAGE fieldをfooterへ追加する。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run_font(run, size_pt=9.0, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _configure_document(document: Document) -> None:
    """US Letter、1 inch余白、presetのstyle tokenを明示する。"""
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after, color in (
        (1, 16, 16, 8, BLUE),
        (2, 13, 12, 6, BLUE),
        (3, 12, 8, 4, DARK_BLUE),
    ):
        style = document.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(header.add_run("SpFlow | Alignment Method Evaluation"), size_pt=9.0, color=MUTED)
    _add_page_number(section.footer.paragraphs[0])


def build_document() -> Path:
    """評価条件、結果、図、限定結論を一つのWord文書へ構成する。"""
    document = Document()
    _configure_document(document)

    # memo_masthead型の冒頭で文書のレビュー目的と版を明示する。
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(title.add_run("Formal S2a/T2a Feasibility Evaluation"), size_pt=24, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run_font(
        subtitle.add_run("Finite residual FIR, real integer-delay buffer, and time-domain comparison"),
        size_pt=13,
        color=MUTED,
    )
    for label, value in (
        ("Scope", "S2a / T2a with completed EBAE weights"),
        ("Main taps", "32 / 128 / 512"),
        ("Signal", "40--88 Hz band-limited broadband, 0 deg / 180 deg endfire"),
        ("Status", "Word review draft before final PDF publication"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _set_run_font(paragraph.add_run(f"{label}: "), size_pt=10.5, bold=True)
        _set_run_font(paragraph.add_run(value), size_pt=10.5)

    _add_heading(document, "1. Executive conclusion", 1)
    _add_body(
        document,
        "T2a preserves the true endfire direction, target level, waveform, SNR, and block continuity after the formal 512-tap realization. "
        "S2a becomes high-rank at the covariance stage and its MUSIC association moves to 42 deg / 138 deg; increasing FIR length cannot recover this failure.",
    )
    _add_bullet(document, "T2a at 512 taps: energy containment 0.992, in-band interior amplitude error 0.69 dB, waveform correlation about 0.997.")
    _add_bullet(document, "T2a estimates three AIC signals instead of the true count of one; formal adoption of the complete EBAE chain is therefore not established.")
    _add_bullet(document, "Do not apply the 512-tap residual-FIR convergence result to the required length of direct-input-coordinate S1/T1 FIRs.")

    _add_heading(document, "2. Evaluation conditions", 1)
    _add_table(
        document,
        ("Item", "Setting"),
        (
            ("Array", "64-channel ULA, 6.25 m spacing, 393.75 m aperture"),
            ("Signal", "Real-time band-limited broadband random signal, 40--88 Hz, 0 / 180 deg"),
            ("Analysis", "fs=1024 Hz, NFFT=16, analysis width 64 Hz"),
            ("Covariance", "4096 non-overlap snapshots, N/E AIC L=M^2=4096"),
            ("Azimuth grid", "0--180 deg in 2 deg steps"),
            ("Main taps", "32 (short failure) / 128 (intermediate) / 512 (converged representative)"),
            ("Runtime", "Real integer-delay buffer, finite residual FIR, 257-sample blocks"),
            ("Input SNR", "High SNR 30.28 dB; fixed-weight stress at 0 dB"),
        ),
        (1900, 7460),
    )
    _add_body(
        document,
        "The occupied band is below the 120 Hz endfire spatial-alias onset. The observed S2a direction error is therefore classified as covariance/MUSIC misassociation, not a grating lobe.",
        bold_lead="Precomputed alias check: ",
    )

    _add_heading(document, "3. Stage-by-stage evaluation", 1)
    _add_heading(document, "3.1 Covariance, AIC, and MUSIC", 2)
    _add_table(
        document,
        ("Source", "Method", "AIC Ns", "MUSIC peak", "Principal / trace", "Rank-1 residual"),
        (
            ("0 deg", "S2a", "32", "42 deg", "0.0912", "0.9428"),
            ("0 deg", "T2a", "3", "0 deg", "0.9976", "0.0011"),
            ("180 deg", "S2a", "32", "138 deg", "0.0911", "0.9429"),
            ("180 deg", "T2a", "3", "180 deg", "0.9976", "0.0011"),
        ),
        (1050, 1050, 900, 1500, 2430, 2430),
    )
    _add_body(
        document,
        "The S2a same-time 16-sample window spans an aperture delay of about 269 samples, so different channels observe different waveform intervals. "
        "T2a extracts the same wavefront interval using candidate-dependent physical delays and preserves a near-rank-1 covariance and the true direction.",
    )

    _add_heading(document, "3.2 Finite residual FIR", 2)
    _add_table(
        document,
        ("Tap", "Energy containment", "Interior amp. error", "Target level 0/180", "Correlation 0/180", "Status"),
        (
            ("32", "0.883", "4.95 dB", "-1.30 / -1.30 dB", "0.978 / 0.977", "Fail"),
            ("128", "0.967", "2.53 dB", "-0.54 / -0.55 dB", "0.990 / 0.990", "Intermediate"),
            ("512", "0.992", "0.69 dB", "-0.41 / -0.37 dB", "0.997 / 0.997", "Pass"),
        ),
        (700, 1300, 1500, 2100, 2100, 1660),
    )
    _add_body(
        document,
        "Acceptance limits are: BL peak error <=2 deg, target-level error <=0.5 dB, interior in-band amplitude error <=1 dB, waveform correlation >=0.99, energy containment >=0.98, and block/monolithic error <=1e-6.",
    )

    _add_heading(document, "3.3 Role of previously evaluated long taps", 2)
    _add_table(
        document,
        ("Prior tap", "Energy containment", "Interior amp. error", "Target level 0/180", "Correlation", "Role"),
        (
            ("256", "0.983", "0.74 dB", "-0.34 / -0.35 dB", "0.995", "Prior minimum pass"),
            ("1024", "0.996", "0.40 dB", "-0.45 / -0.46 dB", "0.998", "Change beyond 512"),
            ("2048", "0.998", "0.17 dB", "+0.01 / -0.01 dB", "0.9986", "Abnormal-case reference"),
        ),
        (900, 1300, 1400, 2100, 1200, 2460),
    )
    _add_body(
        document,
        "The 1024/2048-tap results remain valid historical references. Run them again only if 512 taps fail, 128 to 512 does not converge, energy containment is below 0.98, interior error exceeds 1 dB, or tap shortage cannot be separated from covariance-method differences.",
    )

    _add_heading(document, "3.4 Real-time output", 2)
    _add_table(
        document,
        ("Source", "Method", "512-tap level", "Correlation", "Output SNR", "Status"),
        (
            ("0 deg", "S2a", "-0.90 dB", "0.997", "27.62 dB", "MUSIC mismatch"),
            ("0 deg", "T2a", "-0.41 dB", "0.997", "48.02 dB", "Pass"),
            ("180 deg", "S2a", "-0.86 dB", "0.997", "27.63 dB", "MUSIC mismatch"),
            ("180 deg", "T2a", "-0.37 dB", "0.997", "47.31 dB", "Pass"),
        ),
        (1100, 1000, 1900, 1200, 1700, 2460),
    )
    _add_body(
        document,
        "All 12 primary-sweep cases are finite, and the maximum difference between 257-sample block processing and monolithic processing is zero. "
        "S2a preserves the target waveform if the true direction is selected externally, but it cannot select that direction in the scan and is therefore not judged deployable.",
    )

    _add_heading(document, "4. Visual evidence", 1)
    _add_figure(
        document,
        "abc_summary.png",
        "Figure 1. Stage results from covariance through finite FIR (40--88 Hz broadband; 32/128/512 taps).",
    )
    _add_figure(
        document,
        "bl_broadband_mixed_high_snr.png",
        "Figure 2. High-SNR mixed BL (40--88 Hz broadband; input-band SNR 30.28 dB).",
    )
    _add_figure(
        document,
        "bl_broadband_mixed_low_snr.png",
        "Figure 3. Low-SNR mixed BL (40--88 Hz broadband; 0 dB input-band SNR; fixed-weight stress).",
    )
    _add_figure(
        document,
        "frequency_response.png",
        "Figure 4. In-band frequency response of the finite residual FIR.",
    )
    _add_figure(
        document,
        "waveform_block_boundary.png",
        "Figure 5. Monolithic and block-streaming agreement at 512 taps.",
    )

    # 最終判定節を図の直後へ詰め込まず、レビュー項目を独立ページで読み切れるようにする。
    document.add_page_break()
    _add_heading(document, "5. Validity range and remaining work", 1)
    _add_body(
        document,
        "T2a is conditionally valid for broadband-endfire direction, BL, and runtime realization under this scenario. "
        "AIC source-count overestimation remains, so this conclusion is separate from formal adoption of the complete EBAE chain.",
    )
    for item in (
        "Formal low-SNR redesign of covariance, AIC, and EBAE weights",
        "Different integration times, multiple signals, and nearby unequal-level signals",
        "150 deg, different frequencies and bandwidths, and operationally smooth band edges",
        "Adaptive coefficient-update transients and the validity boundary for AIC overestimation",
    ):
        _add_bullet(document, item)

    note = document.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note_row_properties = note.rows[0]._tr.get_or_add_trPr()
    note_header = OxmlElement("w:tblHeader")
    note_header.set(qn("w:val"), "true")
    note_row_properties.append(note_header)
    _set_cell_shading(note.cell(0, 0), CAUTION_FILL)
    note_paragraph = note.cell(0, 0).paragraphs[0]
    _set_run_font(note_paragraph.add_run("Important: "), size_pt=10.0, bold=True)
    _set_run_font(
        note_paragraph.add_run(
            "The 0 dB plot is a fixed-weight stress test obtained by scaling noise after high-SNR weight design. "
            "It is not a low-SNR redesign of the adaptive weights."
        ),
        size_pt=10.0,
    )
    _set_table_geometry(note, (TABLE_WIDTH_DXA,))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
