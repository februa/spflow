#!/usr/bin/env python3
"""技術Markdownを固定styleの編集可能なWord文書へ変換する。"""

# pyright: reportMissingImports=false, reportUnknownMemberType=false, reportPrivateUsage=false

from __future__ import annotations

import argparse
import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from latex2mathml.converter import convert as latex_to_mathml
except ModuleNotFoundError:
    # 文書runtimeとproject .venvで依存が分かれる場合も、数式自体を欠落させない。
    # 構造化変換を必須にする利用者はCLIの--require-structured-equationsで早期検出できる。
    latex_to_mathml = None

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
# macOS Wordとheadless LibreOfficeの双方へ同じfont fileを渡せる名称を既定にする。
# 別OSでは--fontでNoto Sans CJK JPなど、その環境で実在する日本語fontを明示する。
DEFAULT_FONT = "Arial Unicode MS"
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
HEADER_FILL = "E8EEF5"
CODE_FILL = "F7F8FA"
MATHML_NAMESPACE = "http://www.w3.org/1998/Math/MathML"


@dataclass(frozen=True)
class ConversionOptions:
    """MarkdownからWordへ変換するときの文書設定を表す。

    入力はtitle、subtitle、font名で、出力Word全体のstyleへ適用する。
    Markdown解析や画像生成そのものは責務に含めない。
    """

    title: str | None = None
    subtitle: str | None = None
    font_name: str = DEFAULT_FONT
    require_structured_equations: bool = False


def _set_run_font(
    run: Any,
    font_name: str,
    *,
    size_pt: float,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    """WordとLibreOfficeの両方へASCII・日本語fontを明示する。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def _configure_document(document: Document, font_name: str) -> None:
    """compact reference guide相当のpage geometryとstyleを設定する。"""
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16.0, 18.0, 10.0, HEADING_BLUE),
        2: (13.0, 14.0, 7.0, HEADING_BLUE),
        3: (12.0, 10.0, 5.0, HEADING_DARK_BLUE),
    }
    for level, (size, before, after, color) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = font_name
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _set_cell_margins(cell: Any) -> None:
    """表文字が罫線へ密着しない余白をDXAで設定する。"""
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: tuple[int, ...]) -> None:
    """tblW、tblGrid、全tcWを同じ固定列幅へ揃える。"""
    if not widths_dxa or sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("table column widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    for tag, value in (("w:tblW", TABLE_WIDTH_DXA), ("w:tblInd", TABLE_INDENT_DXA)):
        element = properties.first_child_found_in(tag)
        if element is None:
            element = OxmlElement(tag)
            properties.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            properties = cell._tc.get_or_add_tcPr()
            width = properties.first_child_found_in("w:tcW")
            if width is None:
                width = OxmlElement("w:tcW")
                properties.append(width)
            width.set(qn("w:w"), str(widths_dxa[index]))
            width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _mark_header_row(table: Any) -> None:
    """複数page表で列の意味を保つため先頭行を反復headerにする。"""
    properties = table.rows[0]._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _clean_inline_markdown(text: str) -> str:
    """Word本文へ不要なMarkdown記号を除き、link先は失わず残す。"""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.replace("**", "").replace("__", "").replace("`", "")


def _add_text(document: Document, text: str, font_name: str, *, style: str | None = None) -> None:
    """本文またはlistを一つのWord paragraphとして追加する。"""
    paragraph = document.add_paragraph(style=style)
    _set_run_font(paragraph.add_run(_clean_inline_markdown(text)), font_name, size_pt=11)


def _parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """pipe tableをheaderとdata rowsへ分解する。"""
    parsed = [
        [_clean_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]
        for line in lines
    ]
    if len(parsed) < 2 or not parsed[0]:
        raise ValueError("Markdown table requires a header and at least one row")
    separator = all(set(cell) <= {"-", ":"} and "-" in cell for cell in parsed[1])
    rows = parsed[2:] if separator else parsed[1:]
    if any(len(row) != len(parsed[0]) for row in rows):
        raise ValueError("Markdown table rows must have the same number of columns")
    return parsed[0], rows


def _table_widths(headers: list[str], rows: list[list[str]]) -> tuple[int, ...]:
    """全cellの文字量から説明列へ広い幅を割り当てる。"""
    column_lengths = [len(header) for header in headers]
    for row in rows:
        for index, value in enumerate(row):
            column_lengths[index] = max(column_lengths[index], min(len(value), 36))
    weights = [max(4, value) for value in column_lengths]
    total = sum(weights)
    widths = [max(650, int(TABLE_WIDTH_DXA * value / total)) for value in weights]
    while sum(widths) > TABLE_WIDTH_DXA:
        index = max(range(len(widths)), key=widths.__getitem__)
        widths[index] -= min(10, sum(widths) - TABLE_WIDTH_DXA)
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return tuple(widths)


def _add_table(document: Document, lines: list[str], font_name: str) -> None:
    """Markdown表を固定geometryのWord表へ変換する。"""
    headers, rows = _parse_table(lines)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), HEADER_FILL)
        cell._tc.get_or_add_tcPr().append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(value), font_name, size_pt=8.5, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index == len(values) - 1
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            _set_run_font(paragraph.add_run(value), font_name, size_pt=8.5)
    _set_table_geometry(table, _table_widths(headers, rows))
    _mark_header_row(table)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_record_table(document: Document, lines: list[str], font_name: str) -> None:
    """横長表を行ごとの見出しとlabel/value本文へ変換する。"""
    headers, rows = _parse_table(lines)
    for row in rows:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        _set_run_font(
            heading.add_run(row[0]),
            font_name,
            size_pt=11,
            bold=True,
            color=HEADING_DARK_BLUE,
        )
        for label, value in zip(headers[1:], row[1:]):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            _set_run_font(paragraph.add_run(f"{label}: "), font_name, size_pt=10.5, bold=True)
            _set_run_font(paragraph.add_run(value), font_name, size_pt=10.5)


def _add_image(document: Document, source: Path, alt_text: str, font_name: str) -> None:
    """画像を本文幅内へ挿入し、alt textとcaptionを設定する。"""
    if not source.exists():
        raise FileNotFoundError(f"Markdown image not found: {source}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(source), width=Inches(6.25))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(caption.add_run(alt_text), font_name, size_pt=9, color=MUTED)


def _add_code_block(document: Document, lines: list[str], font_name: str) -> None:
    """改行を保持したcodeを淡色の一列tableへ格納する。"""
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    cell._tc.get_or_add_tcPr().append(shading)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("\n".join(lines))
    _set_run_font(run, "Menlo", size_pt=8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    _set_table_geometry(table, (TABLE_WIDTH_DXA,))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _math_run(text: str) -> Any:
    """Word OMMLへ文字列runを追加する。"""
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    return run


def _append_mathml(parent: Any, element: ET.Element) -> None:
    """MathMLの主要構造を編集可能なWord OMMLへ再帰変換する。"""
    tag = element.tag.removeprefix(f"{{{MATHML_NAMESPACE}}}")
    children = list(element)
    if tag in {"math", "mrow", "semantics"}:
        for child in children:
            if child.tag.endswith("annotation"):
                continue
            _append_mathml(parent, child)
        return
    if tag in {"mi", "mn", "mo", "mtext"}:
        parent.append(_math_run(element.text or ""))
        return
    if tag == "mfrac" and len(children) >= 2:
        fraction = OxmlElement("m:f")
        numerator = OxmlElement("m:num")
        denominator = OxmlElement("m:den")
        _append_mathml(numerator, children[0])
        _append_mathml(denominator, children[1])
        fraction.extend((numerator, denominator))
        parent.append(fraction)
        return
    script_map = {"msub": ("m:sSub", "m:sub"), "msup": ("m:sSup", "m:sup")}
    if tag in script_map and len(children) >= 2:
        outer_name, script_name = script_map[tag]
        outer = OxmlElement(outer_name)
        base = OxmlElement("m:e")
        script = OxmlElement(script_name)
        _append_mathml(base, children[0])
        _append_mathml(script, children[1])
        outer.extend((base, script))
        parent.append(outer)
        return
    if tag == "msubsup" and len(children) >= 3:
        outer = OxmlElement("m:sSubSup")
        base = OxmlElement("m:e")
        subscript = OxmlElement("m:sub")
        superscript = OxmlElement("m:sup")
        _append_mathml(base, children[0])
        _append_mathml(subscript, children[1])
        _append_mathml(superscript, children[2])
        outer.extend((base, subscript, superscript))
        parent.append(outer)
        return
    if tag in {"msqrt", "mroot"} and children:
        radical = OxmlElement("m:rad")
        degree = OxmlElement("m:deg")
        expression = OxmlElement("m:e")
        if tag == "mroot" and len(children) >= 2:
            _append_mathml(degree, children[1])
        _append_mathml(expression, children[0])
        radical.extend((degree, expression))
        parent.append(radical)
        return
    if tag == "mtable":
        matrix = OxmlElement("m:m")
        for row_element in children:
            row = OxmlElement("m:mr")
            for cell_element in list(row_element):
                cell = OxmlElement("m:e")
                _append_mathml(cell, cell_element)
                row.append(cell)
            matrix.append(row)
        parent.append(matrix)
        return
    for child in children:
        _append_mathml(parent, child)


def _add_equation(
    document: Document,
    latex_lines: list[str],
    *,
    require_structured: bool,
) -> None:
    """LaTeX blockをMathML経由でWordネイティブ数式へ変換する。"""
    latex = " ".join(line.strip() for line in latex_lines)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_paragraph = OxmlElement("m:oMathPara")
    equation = OxmlElement("m:oMath")
    if latex_to_mathml is None:
        if require_structured:
            raise RuntimeError(
                "latex2mathml is required when --require-structured-equations is specified"
            )
        equation.append(_math_run(latex))
    else:
        try:
            _append_mathml(equation, ET.fromstring(latex_to_mathml(latex)))
        except (ET.ParseError, ValueError):
            # 未対応演算子でも式を欠落させず、編集可能な文字runとして残す。
            equation.append(_math_run(latex))
    equation_paragraph.append(equation)
    paragraph._p.append(equation_paragraph)


def _add_page_number(paragraph: Any) -> None:
    """footerへWordのPAGE fieldを追加する。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _preflight(lines: list[str]) -> None:
    """意味を保持できないMarkdown構文を変換前に位置付きで拒否する。"""
    in_code = False
    for line_number, line in enumerate(lines, start=1):
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if re.match(r"^\s{2,}[-*+]\s+", line) or re.match(r"^\s{2,}\d+\.\s+", line):
            raise ValueError(f"line {line_number}: nested lists are not supported")
        if stripped.startswith(">"):
            raise ValueError(f"line {line_number}: blockquotes are not supported")
        heading = re.match(r"^(#{5,})\s+", stripped)
        if heading:
            raise ValueError(f"line {line_number}: headings deeper than H4 are not supported")
        if stripped.startswith("<") and stripped.endswith(">") and stripped not in {
            "<!-- pagebreak -->",
            "<!-- word-table: records -->",
        }:
            raise ValueError(f"line {line_number}: unsupported HTML or directive: {stripped}")
    if in_code:
        raise ValueError("unclosed fenced code block")


def convert_markdown(input_path: Path, output_path: Path, options: ConversionOptions) -> Path:
    """MarkdownをWord文書へ変換する。

    Args:
        input_path: UTF-8 Markdown。画像pathはこのfileからの相対path。
        output_path: 生成するDOCX path。
        options: title、subtitle、font設定。

    Returns:
        生成したDOCX path。

    Raises:
        FileNotFoundError: 入力または参照画像が存在しない場合。
        ValueError: 未対応構文、壊れた表、未画像化Mermaidがある場合。
    """
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    lines = input_path.read_text(encoding="utf-8").splitlines()
    _preflight(lines)
    first_h1 = next((line[2:].strip() for line in lines if line.startswith("# ")), input_path.stem)
    title_text = options.title or first_h1

    document = Document()
    _configure_document(document, options.font_name)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run(title_text), options.font_name, size_pt=24, bold=True)
    if options.subtitle:
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(14)
        _set_run_font(
            subtitle.add_run(options.subtitle),
            options.font_name,
            size_pt=13,
            color=MUTED,
        )
    header = document.sections[0].header.paragraphs[0]
    _set_run_font(header.add_run(title_text), options.font_name, size_pt=9, color=MUTED)
    _add_page_number(document.sections[0].footer.paragraphs[0])

    index = 0
    paragraph_buffer: list[str] = []
    table_mode = "table"

    def flush_paragraph() -> None:
        if paragraph_buffer:
            _add_text(document, " ".join(paragraph_buffer), options.font_name)
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped == "<!-- pagebreak -->":
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue
        if stripped == "<!-- word-table: records -->":
            flush_paragraph()
            table_mode = "records"
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            block_start_line = index + 1
            language = stripped[3:].strip().lower()
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            if language == "mermaid":
                raise ValueError(
                    f"line {block_start_line}: Mermaid must be rendered to an image "
                    "before Word conversion"
                )
            _add_code_block(document, block, options.font_name)
            index += 1
            continue
        if stripped == r"\[":
            flush_paragraph()
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != r"\]":
                equation_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError("unclosed display equation")
            _add_equation(
                document,
                equation_lines,
                require_structured=options.require_structured_equations,
            )
            index += 1
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            source_level = len(heading_match.group(1))
            heading_text = _clean_inline_markdown(heading_match.group(2))
            if source_level == 1 and heading_text == first_h1:
                index += 1
                continue
            word_level = max(1, source_level - 1)
            document.add_paragraph(heading_text, style=f"Heading {word_level}")
            index += 1
            continue
        image_match = re.match(r"^!\[([^]]+)\]\(([^)]+)\)$", stripped)
        if image_match:
            flush_paragraph()
            image_path = (input_path.parent / image_match.group(2)).resolve()
            _add_image(document, image_path, image_match.group(1), options.font_name)
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            if table_mode == "records":
                _add_record_table(document, table_lines, options.font_name)
            else:
                _add_table(document, table_lines, options.font_name)
            table_mode = "table"
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            _add_text(document, bullet.group(1), options.font_name, style="List Bullet")
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            _add_text(document, numbered.group(1), options.font_name, style="List Number")
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    """CLI引数を読み、MarkdownからDOCXを生成する。"""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="入力Markdown")
    parser.add_argument("output", type=Path, help="出力DOCX")
    parser.add_argument("--title", help="H1より優先するtitle")
    parser.add_argument("--subtitle", help="title直下のsubtitle")
    parser.add_argument("--font", default=DEFAULT_FONT, help="Word本文font")
    parser.add_argument(
        "--require-structured-equations",
        action="store_true",
        help="latex2mathmlがない場合にLaTeX文字列へfallbackせず終了する",
    )
    arguments = parser.parse_args()
    result = convert_markdown(
        arguments.input,
        arguments.output,
        ConversionOptions(
            title=arguments.title,
            subtitle=arguments.subtitle,
            font_name=arguments.font,
            require_structured_equations=arguments.require_structured_equations,
        ),
    )
    print(result)


if __name__ == "__main__":
    main()
